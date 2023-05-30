import datetime
import os.path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

if TYPE_CHECKING:
    from offer import Offer, User

IMAGE_HEIGHT = Cm(5)

PATH_TO_IMAGES = os.path.join("files", "images")

PATH_TO_COMPANY_INFO = os.path.join("files", "info")

PATH_TO_OFFER = os.path.join("files", "offers")


def specify_font(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)


def specify_margins(document: Document) -> None:
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def create_picture_header(document: Document, logo_filename: str) -> None:
    table = document.add_table(rows=1, cols=1)
    image_cells = table.rows[0].cells
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = IMAGE_HEIGHT

    cell_width = image_cells[0].width
    cell_height = IMAGE_HEIGHT

    paragraph = image_cells[0].paragraphs[0]
    run = paragraph.add_run()

    run.add_picture(
        os.path.join(PATH_TO_IMAGES, logo_filename),
        width=cell_width,
        height=cell_height,
    )

    new_paragraph = document.add_paragraph()
    insert_hr(new_paragraph)


def add_company_info(document: Document):
    with open(os.path.join(PATH_TO_COMPANY_INFO, "requisites.txt")) as file:
        lines = file.readlines()

    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.add_run(line.strip()).bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    insert_hr(document.paragraphs[-1])


def create_offer_header(document: Document, offer: "Offer"):
    date = datetime.date.today()
    string_date = date.strftime("%d.%m.%Y")

    table = document.add_table(rows=2, cols=1)

    offer_info = table.rows[0].cells[0]
    offer_info.text = f"Исх. {offer.number} от {string_date}"
    offer_info.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    with open(os.path.join(PATH_TO_COMPANY_INFO, "offer_header.txt")) as file:
        lines = file.readlines()

    offer_header = table.rows[1].cells[0]
    offer_header.text = ""
    offer_header.paragraphs[0].add_run(lines[0].strip()).bold = True
    offer_header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    offer_header.add_paragraph(lines[1].strip())
    offer_header.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def create_table_header(document: Document, offer: "Offer"):
    table_header = [
        "№",
        "Наименование",
        "Кол-во",
        "Ед.",
        "Срок поставки (дней)",
    ]

    table = document.add_table(rows=1, cols=7, style="Table Grid")
    heading_cells = table.rows[0].cells
    table.columns[0].width = Cm(0.5)
    table.columns[1].width = Cm(10)
    table.columns[2].width = Cm(2)
    table.columns[3].width = Cm(1)

    vat_header = []

    if offer.vat == "10%":
        vat_header.extend(["Цена с НДС 10%", "Сумма с НДС 10%"])
    elif offer.vat == "20%":
        vat_header.extend(["Цена с НДС 20%", "Сумма с НДС 20%"])
    else:
        vat_header.extend(["Цена без НДС", "Сумма без НДС"])

    table_header.extend(vat_header)

    for i in range(len(table_header)):
        heading_cells[i].text = table_header[i]
        heading_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return table


def fill_table(table, offer: "Offer"):
    products = offer.products
    for i in range(len(products)):
        data_row = table.add_row().cells
        data_row[0].text = str(i + 1)
        data_row[1].text = products[i].name
        data_row[2].text = str(products[i].quantity)
        data_row[3].text = products[i].unit
        data_row[4].text = str(products[i].days)
        data_row[5].text = str(products[i].price)
        data_row[6].text = str(products[i].total)

    final_row = table.add_row().cells
    final_row[-2].text = "Итого"
    final_row[-1].text = str(offer.total)


def add_offer_description(document: Document, offer: "Offer"):
    table = document.add_table(rows=1, cols=1)
    description = table.rows[0].cells[0]

    with open(os.path.join(PATH_TO_COMPANY_INFO, "offer_description.txt")) as file:
        lines = file.readlines()

    description.text = ""
    description.paragraphs[0].add_run(lines[0].strip()).bold = True
    description.paragraphs[0].paragraph_format.line_spacing = 1
    description.paragraphs[0].paragraph_format.space_after = Pt(1)

    line_2_start = (
        "Оборудование новое, "
        if offer.supply_type == "Оборудование"
        else "Продукция новая, "
    )
    description.add_paragraph(line_2_start + lines[1].strip())
    description.paragraphs[1].paragraph_format.line_spacing = 1
    description.paragraphs[1].paragraph_format.space_after = Pt(1)

    for line in lines[2:]:
        paragraph = description.add_paragraph(line.strip())
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_after = Pt(1)

    document.add_paragraph()


def add_manager_info(document: Document, user: "User", sign_filename: str):
    with open(os.path.join(PATH_TO_COMPANY_INFO, "ceo_info.txt")) as file:
        lines = file.readlines()
    table = document.add_table(rows=1, cols=3)
    table.columns[0].width = Cm(10)
    table.columns[-1].width = Cm(4)
    manager_info_table = table.rows[0].cells
    manager_info_table[0].text = lines[0]
    manager_info_table[-1].text = lines[-1]

    paragraph = manager_info_table[1].paragraphs[0]
    run = paragraph.add_run()

    run.add_picture(
        os.path.join(PATH_TO_IMAGES, sign_filename), width=Cm(4), height=Cm(4)
    )

    manager_info_cell = manager_info_table[0]

    manager_info_cell.add_paragraph()

    manager_info_cell.add_paragraph(user.position)
    manager_info_cell.add_paragraph(user.full_name)
    manager_info_cell.add_paragraph(user.phone)
    manager_info_cell.add_paragraph(user.email)
    manager_info_cell.add_paragraph(user.website)

    for i in range(-1, -6, -1):
        manager_info_cell.paragraphs[i].paragraph_format.line_spacing = 1
        manager_info_cell.paragraphs[i].paragraph_format.space_after = Pt(1)


def form_docx_offer(
    offer: "Offer",
    user: "User",
    offer_filename: str,
    logo_filename: str,
    sign_filename: str,
):
    document = Document()
    specify_font(document)
    specify_margins(document)
    create_picture_header(document, logo_filename)
    add_company_info(document)
    create_offer_header(document, offer)
    table = create_table_header(document, offer)
    fill_table(table, offer)
    add_offer_description(document, offer)
    add_manager_info(document, user, sign_filename)

    document.save(os.path.join(PATH_TO_OFFER, f"КП-{offer_filename}.docx"))


def insert_hr(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    p_pr.insert_element_before(
        p_bdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
