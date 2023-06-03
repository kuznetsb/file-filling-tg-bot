import os
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm, Inches

from handlers.docx_writer import IMAGE_HEIGHT, insert_hr

if TYPE_CHECKING:
    from invoice import Invoice


PATH_TO_INVOICE = os.path.join("files", "invoices")
PATH_TO_IMAGES = os.path.join("files", "images")
PATH_TO_COMPANY_INFO = os.path.join("files", "info")


def specify_font(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)


def specify_margins(document: Document) -> None:
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def create_picture_header(document: Document, logo_filename: str) -> None:
    table = document.add_table(rows=1, cols=1)
    image_cells = table.rows[0].cells
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Cm(3)

    cell_height = Cm(3)

    paragraph = image_cells[0].paragraphs[0]
    run = paragraph.add_run()

    run.add_picture(
        os.path.join(PATH_TO_IMAGES, logo_filename),
        height=cell_height,
    )


def add_header_info(document: Document):
    with open(os.path.join(PATH_TO_COMPANY_INFO, "invoice_header.txt")) as file:
        lines = file.readlines()

    for i in range(len(lines)):
        if i in (0, (len(lines) - 1)):
            paragraph = document.add_paragraph()
            paragraph.add_run(lines[i].strip()).bold = True
        else:
            paragraph = document.add_paragraph(lines[i].strip())
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_company_info(document: Document):
    with open(os.path.join(PATH_TO_COMPANY_INFO, "invoice_company_info.txt")) as file:
        lines = file.readlines()

    table = document.add_table(rows=4, cols=4, style="Table Grid")

    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(2, 0).merge(table.cell(3, 1))

    table.cell(0, 0).text = lines[1]
    table.cell(0, 1).text = lines[2]
    table.cell(0, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    table.cell(0, 2).text = "Сч. №"
    table.cell(0, 3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    table.cell(0, 3).text = lines[4].split()[-1]
    table.cell(1, 0).text = f"Получатель: {lines[0]}"
    table.cell(2, 0).text = f"Банк получателя: {lines[5]}"
    table.cell(2, 2).text = "БИК"
    table.cell(2, 3).text = lines[6].split()[-1]
    table.cell(3, 2).text = "Сч. №"
    table.cell(3, 3).text = lines[7].split()[-1]


def create_invoice_header(document: Document, invoice: "Invoice"):
    document.add_paragraph()
    with open(os.path.join(PATH_TO_COMPANY_INFO, "invoice_company_info.txt")) as file:
        lines = file.readlines()

    table = document.add_table(rows=4, cols=2)
    table.columns[1].width = Cm(30)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style("HeaderStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = "Times New Roman"

    table.rows[0].cells[1].paragraphs[0].add_run(
        f"СЧЕТ № {invoice.number} от {invoice.date.strftime('%d.%m.%Y')}",
        style="HeaderStyle",
    ).bold = True
    paragraph_row_1 = table.cell(0, 1).add_paragraph()
    insert_hr(paragraph_row_1)
    table.rows[0].cells[1].width = Cm(30)

    table.cell(1, 0).text = "Грузоотправитель"
    table.cell(1, 1).text = (
        f"{lines[0]} ИНН/КПП {lines[1].split()[-1]}/{lines[2].split()[-1]}, "
        f"{lines[3]}"
    )
    paragraph_row_2 = table.cell(1, 1).add_paragraph()
    insert_hr(paragraph_row_2)
    table.rows[1].cells[1].width = Cm(30)

    table.cell(2, 0).text = "Плательщик"
    table.cell(2, 1).text = invoice.buyer
    paragraph_row_3 = table.cell(2, 1).add_paragraph()
    insert_hr(paragraph_row_3)
    table.rows[2].cells[1].width = Cm(30)

    table.cell(3, 0).text = "Грузополучатель"
    table.cell(3, 1).text = invoice.receiver
    paragraph_row_4 = table.cell(3, 1).add_paragraph()
    insert_hr(paragraph_row_4)
    table.rows[3].cells[1].width = Cm(30)


def create_table_header(document: Document, invoice: "Invoice"):
    table_header = [
        "№",
        "Наименование Товара",
        "Цена",
        "Кол-во",
        "Ед.изм",
        "Срок поставки (дней)",
        "Сумма",
    ]

    table = document.add_table(rows=1, cols=7, style="Table Grid")

    heading_cells = table.rows[0].cells
    table.columns[0].width = Cm(0.5)
    table.columns[1].width = Cm(10)
    table.columns[2].width = Cm(2)
    table.columns[3].width = Cm(2)
    table.columns[4].width = Cm(1)

    for i in range(len(table_header)):
        heading_cells[i].text = table_header[i]
        heading_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AUTO

    return table


def fill_table(table, invoice: "Invoice"):
    products = invoice.products
    for i in range(len(products)):
        data_row = table.add_row().cells
        data_row[0].text = str(i + 1)
        data_row[1].text = products[i].name
        data_row[2].text = str(products[i].price)
        data_row[3].text = str(products[i].quantity)
        data_row[4].text = products[i].unit
        data_row[5].text = str(products[i].days)
        data_row[6].text = str(products[i].total)


def add_table_sumup(document, invoice: "Invoice"):
    document.add_paragraph().add_run(f"Итого: {invoice.total}").bold = True
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document.paragraphs[-1].paragraph_format.space_after = Pt(1)

    vat_type = 0

    if invoice.vat == "20%":
        vat_type = 20
    elif invoice.vat == "10%":
        vat_type = 10

    vat = invoice.total / (100 + vat_type) * vat_type

    vat_text = f"в т.ч. НДС {invoice.vat}: {round(vat, 2): }" if vat > 0 else "Без НДС"
    document.add_paragraph().add_run(vat_text).bold = True
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def add_invoice_sumup(document, invoice: "Invoice"):
    table = document.add_table(rows=1, cols=1)
    i, d = divmod(invoice.total, 1)
    digited_total = (
        f"Всего наименований {len(invoice.products)}, "
        f"на сумму {int(i): } руб. {int(d * 100):02d} копеек"
    )
    table.cell(0, 0).text = digited_total
    if invoice.total_text:
        cell = table.add_row().cells[0]
        cell.paragraphs[0].add_run(invoice.total_text.capitalize()).bold = True


def add_ceo_info(document, sign_file):
    with open(os.path.join(PATH_TO_COMPANY_INFO, "ceo_accountant_info.txt")) as file:
        lines = file.readlines()

    table = document.add_table(rows=1, cols=3)
    table.columns[0].width = Cm(10)
    table.columns[-1].width = Cm(4)
    table.cell(0, 0).text = lines[0]
    table.cell(0, 2).text = lines[1]

    paragraph = table.rows[0].cells[1].paragraphs[0]
    run = paragraph.add_run()

    run.add_picture(os.path.join(PATH_TO_IMAGES, sign_file), width=Cm(4), height=Cm(4))

    table.rows[0].cells[0].add_paragraph()
    table.rows[0].cells[0].add_paragraph(lines[2])

    table.rows[0].cells[-1].add_paragraph()
    table.rows[0].cells[-1].add_paragraph(lines[3])


def form_docx_invoice(
    invoice: "Invoice",
    offer_filename: str,
    logo_filename: str,
    sign_filename: str,
):
    document = Document()
    specify_font(document)
    specify_margins(document)
    create_picture_header(document, logo_filename)
    add_header_info(document)
    add_company_info(document)
    create_invoice_header(document, invoice)
    table = create_table_header(document, invoice)
    fill_table(table, invoice)
    add_table_sumup(document, invoice)
    add_invoice_sumup(document, invoice)
    add_ceo_info(document, sign_filename)

    document.save(os.path.join(PATH_TO_INVOICE, f"СЧЕТ-{offer_filename}.docx"))
